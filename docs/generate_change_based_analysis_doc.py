"""Generate change-based-analysis-plan.docx documenting the change-based analysis sub-project."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, color_hex):
    from lxml import etree

    tc_pr = cell._element.get_or_add_tcPr()
    shading_elm = etree.SubElement(tc_pr, qn("w:shd"))
    shading_elm.set(qn("w:val"), "clear")
    shading_elm.set(qn("w:color"), "auto")
    shading_elm.set(qn("w:fill"), color_hex)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2F5496")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def bold_para(doc, label, body):
    p = doc.add_paragraph("")
    run = p.add_run(label + " -- ")
    run.bold = True
    p.add_run(body)


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    title = doc.add_heading(
        "WalmartPricing NLC Pipeline -- Change-Based Analysis", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Context
    doc.add_heading("Context & Motivation", level=1)
    doc.add_paragraph(
        "The existing notebooks/correlation_analysis.ipynb is a 29-cell orchestration notebook "
        "with 14 src/analysis/* modules behind it. An audit (2026-04-24) found that every "
        "headline regression / causal / elasticity analysis in that notebook uses ABSOLUTE levels "
        "of cost_to_walmart, te_margin, and walmart_margin as predictors. The *_vs_7d_pct "
        "percent-change columns are computed in data_prep.py but only one secondary test "
        "(price_change_revenue_analysis) touches cost_to_walmart_vs_7d_pct. The two margin % "
        "change columns (te_margin_vs_7d_pct, walmart_margin_vs_7d_pct) are referenced by zero "
        "downstream analyses."
    )
    doc.add_paragraph(
        "Business question: controlling for context, do CHANGES in price, TE margin, and "
        "Walmart margin (expressed as % vs 7-day average) predict changes in qty_sold and "
        "revenue at the SKU-Node-Date level?"
    )
    doc.add_paragraph(
        "Why a separate notebook: the existing notebook is already large. A standalone notebook "
        "lets the change-based analysis evolve independently without dragging in the "
        "elasticity/DiD/segmented/optimization/strategy machinery."
    )

    # Dataset
    doc.add_heading("Dataset Specification", level=1)

    doc.add_heading("Scope", level=2)
    scope_rows = [
        ("Granularity", "One row per SKU-Node-Date (daily)"),
        ("Analysis window", "90 days (matches correlation_analysis.yaml: data_prep.analysis_days)"),
        ("SKU filter", "Top 90% by qty over 60-day lookback (reused)"),
        ("Rollback exclusion", "Same as the existing notebook"),
        ("Input path", "Fresh DataPreparation.build_master() (5-15 min runtime)"),
        ("Expected size", "~1.3M rows pre-filter; ~30-100K after restricting to change events"),
    ]
    add_styled_table(doc, ["Parameter", "Value"], scope_rows, col_widths=[1.8, 5.2])
    doc.add_paragraph("")

    doc.add_heading("Predictors", level=2)
    pred_rows = [
        ("cost_to_walmart_vs_7d_pct", "(cost - cost_7d_avg) / cost_7d_avg", "data_prep.py:661"),
        ("te_margin_vs_7d_pct", "Same formula applied to te_margin", "data_prep.py:661"),
        ("walmart_margin_vs_7d_pct", "Same formula applied to walmart_margin", "data_prep.py:661"),
    ]
    add_styled_table(
        doc, ["Variable", "Definition", "Source"], pred_rows, col_widths=[2.4, 3.0, 1.8]
    )
    doc.add_paragraph("")

    doc.add_heading("Controls & outcomes", level=2)
    ctrl_rows = [
        ("can_show_inv", "Control", "Binary inventory availability"),
        ("day_of_week", "Control", "Demand seasonality"),
        ("n_active_nodes", "Control", "Network breadth at SKU level"),
        ("days_since_price_change", "Control", "Recency / discount-decay"),
        ("price_tier", "Control (segmentation)", "Pre-existing quartile bucket"),
        ("qty_sold", "Outcome (primary)", "Daily units sold per SKU-Node"),
        ("revenue", "Outcome (primary)", "total_inv_amount per SKU-Node-Date"),
        ("P(qty_sold > 0)", "Outcome (robustness)", "Binary, used as sanity check"),
    ]
    add_styled_table(
        doc, ["Variable", "Role", "Notes"], ctrl_rows, col_widths=[2.2, 1.8, 3.2]
    )
    doc.add_paragraph("")

    # Event definitions
    doc.add_heading("Event-Definition Variants", level=1)
    doc.add_paragraph(
        "Three variant definitions are produced and carried through all relevant analyses. "
        "Sensitivity across variants is the key robustness check."
    )
    variant_rows = [
        ("V1 -- Continuous", "Raw _vs_7d_pct values (includes 0 and small fluctuations)",
         "Regression predictor (OLS, NegBin, elasticity)"),
        ("V2 -- Threshold", "Binary: |_vs_7d_pct| >= 0.01 (matches min_price_change_pct)",
         "DiD treatment flag; ITS event marker"),
        ("V3 -- Directional buckets",
         "5-bin: [-inf,-5%] / (-5%,-1%] / (-1%,1%] / (1%,5%] / (5%,+inf)",
         "Categorical regressor; ANOVA-style group comparison"),
    ]
    add_styled_table(
        doc, ["Variant", "Definition", "Primary use"], variant_rows, col_widths=[1.8, 3.0, 2.4]
    )
    doc.add_paragraph(
        "All three variants are constructed for each of the three change variables -> 9 variant "
        "columns total."
    )

    # Analyses
    doc.add_heading("Analyses", level=1)
    analysis_rows = [
        ("1", "OLS main effects", "3x V1 + controls", "qty_sold",
         "cost + te + wm + controls", "statistical_tests.ols_regression"),
        ("2", "Log-OLS on revenue", "3x V1 + controls", "log(1+revenue)",
         "Same spec, log outcome", "statistical_tests.ols_regression"),
        ("3", "Two-way FE OLS", "3x V1", "qty_sold",
         "SKU-Node + date fixed effects", "Inline (linearmodels.PanelOLS)"),
        ("4", "Negative Binomial", "3x V1", "qty_sold | qty>0",
         "NegBin with log link", "statsmodels.NegativeBinomial"),
        ("5", "Semi-log own-price elasticity",
         "cost_to_walmart_vs_7d_pct", "log(1+qty_sold)",
         "log_qty ~ beta * pct_price", "Adapt elasticity.estimate_elasticity"),
        ("6", "FE panel elasticity",
         "V1 price only", "log(1+qty_sold)",
         "Two-way FE semi-log", "Adapt elasticity.estimate_elasticity_fe"),
        ("7", "DiD on threshold events",
         "V2 price-threshold", "qty_sold, revenue",
         "Treated = first V2 event; same-brand controls",
         "did_effects.build_did_panel + heterogeneous_did"),
        ("8", "Interrupted time series",
         "V2 event markers", "qty_sold",
         "Segmented regression, top 20 SKU-Nodes, HAC SEs", "Inline ITS helper"),
        ("9", "Directional bucket ANOVA",
         "V3 buckets", "qty_sold, revenue",
         "One-way ANOVA + Tukey HSD per variable", "Inline scipy + statsmodels"),
        ("10", "Interaction sensitivity (one-off)",
         "price_pct x te_pct, price_pct x wm_pct", "qty_sold",
         "Single OLS, main effects retained", "statistical_tests.ols_regression extended"),
    ]
    add_styled_table(
        doc,
        ["#", "Analysis", "Predictors", "Outcome", "Spec", "Reuses"],
        analysis_rows,
        col_widths=[0.3, 1.5, 1.2, 1.0, 1.8, 1.5],
    )
    doc.add_paragraph(
        "Deliberately excluded (keeping notebook tight): logistic P(sale>0), propensity matching, "
        "clustering, Lorenz/Gini, optimization, strategy recommendations -- these remain in the "
        "main correlation notebook if needed."
    )

    # Notebook structure
    doc.add_heading("Notebook Structure (~18 cells)", level=1)
    nb_rows = [
        ("1. Setup", "1-2", "Imports, config, path setup"),
        ("2. Data load", "3-4",
         "DataPreparation.build_master(); sanity-check three _vs_7d_pct columns"),
        ("3. Variant construction", "5",
         "Build V2 threshold flags and V3 directional buckets for the three change variables"),
        ("4. Sample characterization", "6",
         "Change vs no-change counts; distributions; bucket sizes; VIF table"),
        ("5. OLS suite", "7-8", "Analyses 1, 2, 3 -- qty_sold and revenue, pooled and FE"),
        ("6. NegBin on sale-days", "9", "Analysis 4"),
        ("7. Price elasticity", "10-11", "Analyses 5, 6 -- semi-log + FE panel"),
        ("8. DiD", "12", "Analysis 7 -- threshold treatment on cost_to_walmart_vs_7d_pct"),
        ("9. ITS", "13", "Analysis 8 -- top-20 SKU-Node segmented regressions"),
        ("10. Directional buckets", "14-15",
         "Analysis 9 -- ANOVA + pairwise comparisons for each change var"),
        ("11. Interaction robustness", "16", "Analysis 10 -- one-off sensitivity model"),
        ("12. Summary", "17",
         "Coefficient forest plot across all specs; findings markdown; export"),
        ("13. Cleanup", "18", "loader.close(); export summary parquet"),
    ]
    add_styled_table(
        doc, ["Section", "Cells", "Purpose"], nb_rows, col_widths=[2.0, 0.8, 4.4]
    )
    doc.add_paragraph(
        "Summary export: outputs/change_based_coefficients_{end_date}.parquet -- one row per "
        "(analysis, variant, predictor, outcome) with coef, se, pvalue, ci_low, ci_high, "
        "n_obs, r2_or_pseudo."
    )

    # Config
    doc.add_heading("Config Additions", level=1)
    doc.add_paragraph(
        "New change_based: section added to config/correlation_analysis.yaml. No new "
        "data-loading config keys -- inherits from data_prep block."
    )
    cfg_rows = [
        ("change_cols", "3 _vs_7d_pct columns"),
        ("outcome_cols", "qty_sold, revenue"),
        ("controls", "can_show_inv, day_of_week, n_active_nodes, days_since_price_change"),
        ("variants.threshold_pct", "0.01 (matches pipeline min_price_change_pct)"),
        ("variants.directional_bins",
         "[-1.0, -0.05, -0.01, 0.01, 0.05, 1.0] (5-bin edges)"),
        ("its.top_n_sku_nodes", "20"),
        ("its.min_events", "2"),
        ("its.pre_window / post_window", "7 / 14 days"),
        ("did.treatment_col", "cost_to_walmart_vs_7d_pct"),
        ("did.threshold", "0.01"),
        ("min_obs_nb", "200 (NegBin fit guard)"),
        ("interactions", "[(cost_pct, te_pct), (cost_pct, wm_pct)]"),
    ]
    add_styled_table(doc, ["Config key", "Value"], cfg_rows, col_widths=[2.8, 4.4])
    doc.add_paragraph("")

    # Reuse strategy
    doc.add_heading("Module Reuse & New Code", level=1)

    doc.add_heading("Reused as-is (via config override)", level=2)
    bold_para(
        doc,
        "data_prep.DataPreparation.build_master",
        "Called unchanged to assemble the master df.",
    )
    bold_para(
        doc,
        "statistical_tests.ols_regression",
        "Called with feature_cols swapped to the 3 _vs_7d_pct predictors + controls.",
    )
    bold_para(doc, "ci_utils.*", "All CI helpers used unchanged.")

    doc.add_heading("Adapted via thin wrappers (inline in notebook)", level=2)
    wrap_rows = [
        ("Semi-log elasticity (analysis 5)",
         "20-line inline wrapper: log1p(qty) ~ pct via statsmodels OLS; beta, SE, 95% CI"),
        ("FE panel semi-log (analysis 6)",
         "30-line inline wrapper around linearmodels.PanelOLS with two-way FE"),
        ("DiD with % threshold treatment (analysis 7)",
         "Call did_effects.build_did_panel with treatment_col = cost_to_walmart_vs_7d_pct"),
        ("NegBin (analysis 4)", "15-line inline statsmodels call"),
        ("ITS (analysis 8)",
         "40-line inline function: segmented regression pre/post V2 events with HAC SEs"),
        ("ANOVA + Tukey (analysis 9)",
         "20-line scipy + statsmodels wrapper iterating over V3 buckets"),
    ]
    add_styled_table(doc, ["What", "How"], wrap_rows, col_widths=[2.4, 4.8])
    doc.add_paragraph("")

    doc.add_heading("No new modules under src/analysis/", level=2)
    doc.add_paragraph(
        "Deliberate choice: the notebook stays self-contained. If any inline helper proves "
        "reusable (likely ITS), it graduates to src/analysis/change_based.py in a follow-up -- "
        "not as part of this plan."
    )

    # Caveats
    doc.add_heading("Caveats & Decisions", level=1)
    caveat_rows = [
        ("Log-log elasticity fails with signed % changes",
         "Use semi-log (log(1+qty) ~ pct_price) -- coefficient is % change in qty per 1pp in price"),
        ("Zero / NaN denominator in _vs_7d_pct",
         "Handled in data_prep.py:663 (replace(0, nan)); rows with NaN predictor dropped per-analysis"),
        ("Most SKU-Node-days have zero % change",
         "V1 keeps all rows; V2 / V3 filter or bucket explicitly"),
        ("Rolling window induces autocorrelation in predictor",
         "HC1 SEs in OLS; HAC (Newey-West) in ITS; clustered SEs in FE panel"),
        ("Multicollinearity between cost_pct, te_pct, wm_pct",
         "VIF in characterization cell; if > 10, run single-predictor variants as robustness"),
        ("Large zero-sales mass",
         "NegBin restricts to qty > 0; bucket ANOVA compares means regardless of sparsity"),
        ("Interpretation of interaction terms",
         "Analysis 10 marked robustness-only; headline conclusions rest on main-effect models"),
    ]
    add_styled_table(doc, ["Issue", "Decision"], caveat_rows, col_widths=[2.8, 4.4])
    doc.add_paragraph("")

    # Checklist
    doc.add_heading("Verification Checklist", level=1)
    checklist = [
        "Notebook runs end-to-end with default 90-day window in < 25 min",
        "All three _vs_7d_pct columns exist in master df with > 50% non-NaN",
        "VIF reported in characterization cell; all < 10 for chosen main-effect spec",
        "Coefficient signs across V1, V2, V3 point in the same direction for price %",
        "DiD parallel-trends plot rendered; visual inspection passes",
        "ITS runs successfully on >= 15 of top-20 SKU-Nodes",
        "Summary parquet written with >= one row per analysis",
        "No import of src/analysis/* beyond data_prep, statistical_tests, did_effects, ci_utils",
        "Dependencies linearmodels (FE panel) and statsmodels (NegBin) verified at notebook start",
    ]
    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")

    # Deliverables
    doc.add_heading("Deliverables", level=1)
    deliv_rows = [
        ("docs/change-based-analysis-plan.md", "This document (markdown)"),
        ("docs/change-based-analysis-plan.docx", "Generated from generate_change_based_analysis_doc.py"),
        ("notebooks/change_based_analysis.ipynb", "To build after plan approval"),
        ("config/correlation_analysis.yaml",
         "Add change_based: section (non-breaking)"),
        ("outputs/change_based_coefficients_{date}.parquet", "Generated on notebook run"),
    ]
    add_styled_table(doc, ["File", "Status"], deliv_rows, col_widths=[3.0, 4.2])
    doc.add_paragraph("")

    # Out of scope
    doc.add_heading("Out of Scope (for this sub-project)", level=1)
    oos = [
        "Absolute-level analyses (stay in the main correlation notebook)",
        "Logistic regression for P(sale>0) (already in main notebook)",
        "A/B tests for Wm margin split / brand margin arms (already in main notebook)",
        "Propensity score matching, clustering, node efficiency (already in main notebook)",
        "Writing change-based results back into pricing rules / simulation / strategy modules",
    ]
    for item in oos:
        doc.add_paragraph(item, style="List Bullet")

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Generated 2026-04-24 -- WalmartPricing NLC Pipeline -- Change-Based Analysis Plan"
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    return doc


def main():
    doc = build_document()
    out_path = Path(__file__).resolve().parent / "change-based-analysis-plan.docx"
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()

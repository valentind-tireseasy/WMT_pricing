"""Generate PROJECT-PLAN.docx from the project plan content."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from lxml import etree

    tc_pr = cell._element.get_or_add_tcPr()
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    shading_elm = etree.SubElement(tc_pr, qn("w:shd"))
    shading_elm.set(qn("w:val"), "clear")
    shading_elm.set(qn("w:color"), "auto")
    shading_elm.set(qn("w:fill"), color_hex)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2F5496")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def build_document():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    # ── Title ──
    title = doc.add_heading("WalmartPricing NLC Pipeline — Project Plan", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Context ──
    doc.add_heading("Context", level=1)
    doc.add_paragraph(
        "The Walmart B2B Node Level Cost (NLC) pricing system manages pricing for ~100K SKUs "
        "across warehouse nodes on the Walmart B2B platform. Unlike Amazon (price per SKU, 4 regions, "
        "Contra COGs Excel upload), Walmart uses per-SKU-Node pricing — each warehouse gets its own "
        "cost, delivered as a DSV CSV upload via the hybris platform."
    )
    doc.add_paragraph(
        "The system was previously a monolithic 242-cell Jupyter notebook "
        "(original_code/node_level_costs_pricing.ipynb) that was run semi-automatically: the notebook "
        "could be executed end-to-end, but required manual parameter toggling, file management, and "
        "hybris upload steps between runs. There was no notification system, no automated validation, "
        "and no config-driven parameterization — all values were hardcoded in notebook cells."
    )
    p = doc.add_paragraph("")
    run = p.add_run(
        "Goal: Transform the notebook into a modular, config-driven Python pipeline with automated "
        "upload, validation, notifications, and diagnostic tools — while preserving the exact NLC "
        "computation logic. Then extend with A/B testing infrastructure to experiment with margin "
        "parameters and track performance."
    )
    run.bold = True

    # ── Prior State ──
    doc.add_heading("Prior State — Original Notebook", level=1)
    doc.add_paragraph(
        "The original system (original_code/node_level_costs_pricing.ipynb) was a single notebook with:"
    )
    items = [
        "242 cells (194 code, 48 markdown) containing all logic inline",
        "Hardcoded file paths, Google Drive IDs, DW table names, and margin thresholds",
        "Manual toggling of optional sections (rollbacks, national prices, tests) via (no) labels in section headers",
        "No separation between data loading, model computation, business rules, and output generation",
        "No error handling or notifications — failures required manual inspection",
        "hybris upload done manually through the browser",
        "FTP validation run manually ~3 hours after upload",
        "Tests tracker updated in-notebook with no backup mechanism",
        "Shared module imports (pricing_module, hybris) mixed into cell execution order",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "The notebook was functional but fragile: a mis-ordered cell execution, a wrong parameter, "
        "or a silent data issue could cascade through the entire run without detection."
    )

    # ── Architecture ──
    doc.add_heading("Architecture — What Was Built", level=1)

    arch_rows = [
        ("config/", "4 YAML config files", "All parameters externalized"),
        ("src/adapters/", "module_loader, google_api_adapter, dw_adapter", "Thin wrappers for external data access"),
        ("src/data/", "loader, inventory_checker", "Config-driven data dispatch + diagnostics"),
        ("src/models/", "nlc_model, run_model", "Core NLC model (two-pass, 8 margins, cascade)"),
        ("src/rules/", "pricing_rules", "5-rule pricing engine"),
        ("src/dsv/", "dsv_builder, hybris_uploader, ftp_validator", "DSV construction, upload, and validation"),
        ("src/tracker/", "tracker_updater", "Tests tracker maintenance"),
        ("src/notifications/", "slack_notifier", "Step-by-step Slack notifications"),
        ("src/pipeline.py", "run_pipeline, run_ftp_validation", "Main orchestrator"),
        ("notebooks/", "01_nlc_pricing.ipynb", "Interactive orchestrator (mirrors pipeline.py)"),
        ("original_code/", "node_level_costs_pricing.ipynb", "Reference only — original 242-cell notebook"),
    ]
    add_styled_table(
        doc,
        ["Directory", "Key Modules", "Purpose"],
        arch_rows,
        col_widths=[1.8, 2.8, 2.6],
    )
    doc.add_paragraph("")

    # ── Phase 1 ──
    doc.add_heading("Phase 1 — Core Pipeline Extraction", level=1)
    p = doc.add_paragraph("")
    run = p.add_run("STATUS: DONE")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    p.add_run("  •  Commits: 1277cba through 2d3e787  •  2026-03-18")

    doc.add_paragraph(
        "Extracted the notebook's logic into modular Python with clean separation of concerns."
    )

    doc.add_heading("What was done", level=2)

    phase1_items = [
        (
            "Config externalization",
            "Moved all hardcoded values (Google Drive IDs, DW table names, margin thresholds, "
            "inventory parameters, file paths) into 4 YAML config files. Every parameter that was "
            "previously buried in a notebook cell is now in config/.",
        ),
        (
            "Adapter layer",
            "Created src/adapters/ to abstract data access. Google Sheets, Google Drive, and Data "
            "Warehouse calls go through thin adapter classes. Shared modules (pricing_module, hybris) "
            "use lazy imports to avoid triggering Selenium/OAuth at startup. The {drive} placeholder "
            "in paths allows the pipeline to run on any machine without path editing.",
        ),
        (
            "Data loader",
            "Built src/data/loader.py as a config-driven dispatcher. Given a source name from "
            "data_sources.yaml, it routes to the correct adapter (Google API, DW query, or local "
            "file read), applies column renames and dtype specs, and returns a clean DataFrame.",
        ),
        (
            "NLC model",
            "Ported the two-pass inventory computation into src/models/nlc_model.py. Preserves the "
            "exact algorithm: filter inventory by min_units threshold and zero-out threshold, compute "
            "NLC at 8 margin levels (6–20%), apply MAP and Walmart margin constraints, cascade to "
            "final NLC (11% → 8% → 6% → N/A). The two-pass strategy (min_units=8 first, then "
            "min_units=4 for gaps) is preserved exactly.",
        ),
        (
            "Pricing rules engine",
            "Extracted the 5 update-type categorization into src/rules/pricing_rules.py: Walmart "
            "margin split test, brand margin test, low price updates, high price updates, new "
            "SKU-Nodes. Protected target exclusion and 1% minimum price change threshold are configurable.",
        ),
        (
            "DSV builder",
            "src/dsv/dsv_builder.py handles DSV construction: starts from current DSV, optionally "
            "applies national price overrides and rollback handling, merges pricing rule updates, "
            "validates the result.",
        ),
        (
            "Tracker updater",
            "src/tracker/tracker_updater.py refreshes margin columns, appends new/updated entries, "
            "deduplicates by (Product Code, Identifier), and creates dated backups on the shared "
            "drive before overwriting.",
        ),
    ]
    for title_text, desc in phase1_items:
        p = doc.add_paragraph("")
        run = p.add_run(title_text + " — ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("Key decisions", level=2)
    decisions = [
        "Preserve algorithm exactly — No optimization or logic changes during extraction. The model produces the same output as the notebook.",
        "Config over code — Business parameters live in YAML, not Python. Changes to thresholds, margin levels, or data sources don't require code changes.",
        "Lazy imports — hybris and pricing_module are imported only when their steps run, avoiding Selenium browser launch and Google OAuth flows on import.",
        "0.5s Google API sleep — Rate limiting between Google API calls during batch operations to avoid quota exhaustion.",
    ]
    for d in decisions:
        doc.add_paragraph(d, style="List Bullet")

    # ── Phase 2 ──
    doc.add_heading("Phase 2 — Orchestration and Interactive Notebook", level=1)
    p = doc.add_paragraph("")
    run = p.add_run("STATUS: DONE")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    p.add_run("  •  Commits: 2b0a7c5 through bcc327e  •  2026-03-18")

    doc.add_heading("What was done", level=2)

    phase2_items = [
        (
            "Pipeline orchestrator",
            "Created src/pipeline.py with run_pipeline(**kwargs) that executes all steps in sequence "
            "with parameter-driven toggling of optional steps. Also exposes run_ftp_validation(today_str) "
            "for the separate post-upload validation run.",
        ),
        (
            "Interactive notebook",
            "Built notebooks/01_nlc_pricing.ipynb that mirrors pipeline.py with parameter cells at the "
            "top. Allows toggling individual steps and running interactively while using the same "
            "underlying modules.",
        ),
        (
            "hybris upload automation",
            "src/dsv/hybris_uploader.py automates the previously manual browser workflow: sign in, "
            "navigate to DSV prices, select the WalmartB2B channel, upload the CSV, and poll for completion.",
        ),
        (
            "FTP validation",
            "src/dsv/ftp_validator.py downloads XML response files from Walmart's FTP, parses ingestion "
            "status per line item, generates a summary Excel, and flags failure rates above 1.5%.",
        ),
        (
            "Local output mode",
            "Added local_output flag to save all files to a local outputs/ directory instead of the "
            "shared drive, enabling safe test runs without overwriting production files.",
        ),
        (
            "First full run",
            "Executed the notebook end-to-end and saved cell outputs, validating the pipeline produces "
            "correct results against the notebook's original output.",
        ),
    ]
    for title_text, desc in phase2_items:
        p = doc.add_paragraph("")
        run = p.add_run(title_text + " — ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("Key decisions", level=2)
    decisions2 = [
        "Two entry points — Both pipeline.py (programmatic) and the notebook (interactive) use the same underlying modules. The notebook is for exploration and debugging; the pipeline is for production runs.",
        "Optional steps via kwargs — Rollbacks, national prices, hybris upload, and inventory check are all toggled by boolean parameters with sensible defaults.",
    ]
    for d in decisions2:
        doc.add_paragraph(d, style="List Bullet")

    # ── Phase 3 ──
    doc.add_heading("Phase 3 — Diagnostics and Configurability", level=1)
    p = doc.add_paragraph("")
    run = p.add_run("STATUS: DONE")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    p.add_run("  •  Commits: d8f21d0 through 5a4d251  •  2026-03-19")

    doc.add_heading("What was done", level=2)
    phase3_items = [
        (
            "Inventory checker",
            "Added src/data/inventory_checker.py to compare today's inventory costs against the previous "
            "run. Computes per-SKU-Warehouse cost deltas and summarizes by brand and vendor (increases "
            "and decreases), giving visibility into cost shifts that drive price changes.",
        ),
        (
            "Configurable shared drive letter",
            "Parameterized the shared drive letter in config/settings.yaml (shared_drive_letter: \"G:\"), "
            "with {drive} placeholder resolution throughout all path templates. The pipeline can run on "
            "any machine by changing one config value.",
        ),
        (
            "Process overview documentation",
            "Created docs/process_overview.md with detailed step-by-step documentation of the pipeline, "
            "parameters, data sources, and business rules.",
        ),
    ]
    for title_text, desc in phase3_items:
        p = doc.add_paragraph("")
        run = p.add_run(title_text + " — ")
        run.bold = True
        p.add_run(desc)

    # ── Phase 4 ──
    doc.add_heading("Phase 4 — Notifications and Production Polish", level=1)
    p = doc.add_paragraph("")
    run = p.add_run("STATUS: DONE")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    p.add_run("  •  Commits: 083defd through e38ac9f  •  2026-03-20 to 2026-03-23")

    doc.add_heading("What was done", level=2)
    phase4_items = [
        (
            "Slack notifications",
            "Built src/notifications/slack_notifier.py with step-by-step Slack posts for the entire "
            "pipeline lifecycle: start, inventory check, NLC model, pricing rules, national prices, "
            "rollbacks, DSV build, tracker update, save, hybris upload, FTP validation, completion, "
            "and errors. Disabled steps are silently skipped (no \"skipped\" clutter).",
        ),
        (
            "4-table inventory breakdown",
            "Slack inventory check notifications show 4 formatted tables: top 5 brand increases, "
            "top 5 brand decreases, top 5 vendor increases, top 5 vendor decreases. Provides immediate "
            "visibility into what's driving price changes.",
        ),
        (
            "Smart inventory date comparison",
            "The inventory checker now compares against the actual last run date (read from last_run.txt) "
            "instead of always comparing to yesterday. Handles gaps between runs correctly.",
        ),
        (
            "hybris polling tuned",
            "Changed polling interval from 30 seconds to 5 minutes to avoid disrupting Walmart's "
            "processing pipeline. Timeout remains at 1 hour.",
        ),
        (
            "DSV archive",
            "After successful hybris upload, the DSV file is automatically copied to the shared drive "
            "archive folder.",
        ),
    ]
    for title_text, desc in phase4_items:
        p = doc.add_paragraph("")
        run = p.add_run(title_text + " — ")
        run.bold = True
        p.add_run(desc)

    # ── Current State Summary ──
    doc.add_heading("Current State — Summary", level=1)

    doc.add_heading("What the pipeline does today", level=2)
    today_items = [
        "Loads 8 data sources from Google Drive, Data Warehouse, and shared drive via config-driven adapters",
        "Computes optimal NLC prices using a two-pass inventory model with 8 margin levels and cascading logic",
        "Applies 5 pricing rules (margin split test, margin test, low/high price updates, new nodes)",
        "Generates DSV files (~3.3M rows) for Walmart upload",
        "Automates hybris upload via Selenium with 5-minute polling",
        "Validates FTP responses ~3 hours post-upload, alerts on high failure rates",
        "Updates tests tracker with per-SKU-Node margin history and dated backups",
        "Posts Slack notifications at each step with detailed breakdowns",
        "Handles optional workflows: rollbacks, national price overrides, inventory diagnostics",
    ]
    for i, item in enumerate(today_items, 1):
        doc.add_paragraph(f"{i}. {item}")

    doc.add_heading("What changed from the original notebook", level=2)

    comparison_rows = [
        ("Structure", "242 cells, all inline", "22 modules across 7 packages"),
        ("Configuration", "Hardcoded in cells", "4 YAML config files"),
        ("Data access", "Direct API calls in cells", "Adapter layer with config dispatch"),
        ("Optional steps", 'Toggle by (no) in headers', "Boolean kwargs with defaults"),
        ("hybris upload", "Manual browser workflow", "Automated Selenium with polling"),
        ("FTP validation", "Manual, run separately", "Automated with failure alerting"),
        ("Notifications", "None", "Step-by-step Slack with breakdowns"),
        ("Inventory diagnostics", "None", "Day-over-day cost comparison"),
        ("Tracker backup", "None", "Dated backup before overwrite"),
        ("Output flexibility", "Shared drive only", "Local or shared drive"),
        ("Path portability", "Hardcoded drive letters", "{drive} placeholder resolution"),
        ("Error visibility", "Silent failures", "Slack error notifications"),
    ]
    add_styled_table(
        doc,
        ["Aspect", "Original Notebook", "Current Pipeline"],
        comparison_rows,
        col_widths=[1.5, 2.5, 3.0],
    )
    doc.add_paragraph("")

    # ── Phase 5 ──
    doc.add_heading("Phase 5 — A/B Testing and Parameter Experimentation", level=1)
    p = doc.add_paragraph("")
    run = p.add_run("STATUS: PLANNED")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)

    doc.add_heading("Objective", level=2)
    doc.add_paragraph(
        "Build infrastructure to systematically experiment with NLC pricing parameters — margin "
        "levels, Walmart margin split percentages, and other model parameters — across defined "
        "groups of SKU-Nodes, then measure and compare performance over time."
    )

    doc.add_heading("Current A/B testing capabilities", level=2)
    doc.add_paragraph("The pipeline already supports two test types:")
    doc.add_paragraph(
        "Walmart margin split test — 3 sub-groups (60% split, 50% split, baseline) applied to "
        "SKU-Nodes tagged \"Wm margin split test\" in the tracker",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Brand margin test — 5 sub-groups (11%–15% margin levels) applied to SKU-Nodes tagged "
        "\"Margin test\", filterable by start date",
        style="List Bullet",
    )
    doc.add_paragraph(
        "These are functional but static: the groups and parameters were defined once and haven't "
        "been iterated on."
    )

    doc.add_heading("Planned work", level=2)

    p = doc.add_paragraph("")
    run = p.add_run("5.1 — New experimental groups")
    run.bold = True
    phase5_1 = [
        "Define new test groups beyond the existing two, targeting specific SKU segments (by brand, sales category, node, or margin band)",
        "Assign experimental margin levels, split percentages, or other parameters to each group",
        "Ensure proper control groups (baseline) for each experiment",
    ]
    for item in phase5_1:
        doc.add_paragraph(item, style="List Bullet")

    p = doc.add_paragraph("")
    run = p.add_run("5.2 — Parameter tuning on existing tests")
    run.bold = True
    phase5_2 = [
        "Adjust parameters on the existing Wm margin split test (e.g., try 55% or 65% split instead of 60%/50%)",
        "Adjust margin levels on the Brand margin test sub-groups",
        "Track when parameter changes were made for before/after comparison",
    ]
    for item in phase5_2:
        doc.add_paragraph(item, style="List Bullet")

    p = doc.add_paragraph("")
    run = p.add_run("5.3 — Performance analysis and reporting")
    run.bold = True
    phase5_3 = [
        "Build analysis to compare experimental vs. control group performance over time",
        "Key metrics: revenue, margin, conversion rate, inventory sell-through",
        "Time-series tracking to account for seasonality and market shifts",
        "Automated reporting (Slack or Excel) on experiment outcomes",
    ]
    for item in phase5_3:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Open questions", level=2)
    questions = [
        "What SKU segments should be targeted for new experiments?",
        "How long should each experiment run before drawing conclusions?",
        "What metrics define \"better performance\" for a given experiment?",
        "Should experiments be managed in the tracker CSV or a separate experiment config?",
    ]
    for q in questions:
        doc.add_paragraph(q, style="List Bullet")

    # ── Git History ──
    doc.add_heading("Git History Reference", level=1)

    git_rows = [
        ("2026-03-18", "1277cba", "Phase 1", "Initial scaffold: config, adapters, model, rules, DSV, tracker"),
        ("2026-03-18", "2d3e787", "Phase 1", "Refine model logic and data sources to match original notebook"),
        ("2026-03-18", "2b0a7c5", "Phase 2", "Add orchestrator notebook with toggleable optional steps"),
        ("2026-03-18", "55e9c98", "Phase 2", "Add hybris DSV upload automation"),
        ("2026-03-18", "70e3c41", "Phase 2", "Add .gitignore and remove cached pycache files"),
        ("2026-03-18", "8d8a24c", "Phase 2", "Fix rollbacks loading: use rollbacks_path parameter"),
        ("2026-03-18", "9cb9b93", "Phase 2", "Add local output mode for safe test runs"),
        ("2026-03-18", "bcc327e", "Phase 2", "First full end-to-end run with saved outputs"),
        ("2026-03-19", "d8f21d0", "Phase 3", "Add pipeline process overview documentation"),
        ("2026-03-19", "b60670d", "Phase 3", "Add inventory check step and configurable shared drive letter"),
        ("2026-03-19", "5a4d251", "Phase 3", "Fran_1"),
        ("2026-03-20", "083defd", "Phase 4", "Add Slack notifications and silence skipped steps"),
        ("2026-03-23", "e38ac9f", "Phase 4", "hybris 5min poll, smart inventory date, 4-table Slack, DSV archive"),
    ]
    add_styled_table(
        doc,
        ["Date", "Commit", "Phase", "Description"],
        git_rows,
        col_widths=[1.1, 0.9, 0.8, 4.4],
    )

    # Save
    output_path = "docs/PROJECT-PLAN.docx"
    doc.save(output_path)
    print(f"Saved to {output_path}")


if __name__ == "__main__":
    build_document()

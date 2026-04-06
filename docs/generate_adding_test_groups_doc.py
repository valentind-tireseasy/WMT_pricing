"""Generate ADDING-TEST-GROUPS.docx from the guide content."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from lxml import etree

    tc_pr = cell._element.get_or_add_tcPr()
    shading_elm = etree.SubElement(tc_pr, qn("w:shd"))
    shading_elm.set(qn("w:val"), "clear")
    shading_elm.set(qn("w:color"), "auto")
    shading_elm.set(qn("w:fill"), color_hex)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2F5496")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def add_code_block(doc, code_text):
    """Add a code block with monospace font and gray background."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Apply shading to the paragraph
    from lxml import etree
    pPr = p._element.get_or_add_pPr()
    shading_elm = etree.SubElement(pPr, qn("w:shd"))
    shading_elm.set(qn("w:val"), "clear")
    shading_elm.set(qn("w:color"), "auto")
    shading_elm.set(qn("w:fill"), "F2F2F2")


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    # ── Title ──
    title = doc.add_heading(
        "How to Add New Test Groups / Recurrent Update Types", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Purpose ──
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This guide explains how to introduce a new pricing test group (like "
        '"DSVD test" or "Increase test") into the WalmartPricing pipeline. It is '
        "written for a technical user who works with the live notebook and needs "
        "to mirror those changes in the modular project code."
    )
    doc.add_paragraph(
        "When a new test group is added in the notebook, 4 files in the project "
        "need to be updated. This document walks through each one."
    )

    # ── Overview ──
    doc.add_heading("Overview: What makes a test group work", level=1)
    doc.add_paragraph(
        'A test group is a label stored in the "Final target" column of the tests '
        'tracker (e.g., "Margin test", "DSVD test", "Increase test"). Each pipeline run:'
    )
    overview_items = [
        "Loads the tracker and merges Final target + Sub-group onto the NLC model output",
        "Applies recurrent update logic for each group (computes new prices based on group-specific rules)",
        "Excludes those tagged SKU-Nodes from the regular low/high price updates",
        "Includes the group's DSV updates in the final DSV file",
        "Includes the group's tracker updates in the tracker save",
    ]
    for i, item in enumerate(overview_items, 1):
        doc.add_paragraph(f"{i}. {item}")

    p = doc.add_paragraph("")
    run = p.add_run("To add a new group, you touch these 4 files:")
    run.bold = True

    file_rows = [
        ("1", "config/nlc_model.yaml", "Add the group name to dont_update_targets"),
        ("2", "src/rules/pricing_rules.py", "Add a new method that computes prices for the group"),
        ("3", "src/pipeline.py", "Call the new method, wire outputs into DSV + tracker"),
        ("4", "config/pricing_rules.yaml", "(Optional) Add config entries for configurable parameters"),
    ]
    add_styled_table(
        doc,
        ["#", "File", "What to change"],
        file_rows,
        col_widths=[0.4, 2.5, 4.3],
    )
    doc.add_paragraph("")

    # ── Step 1 ──
    doc.add_heading("Step 1: Add to the exclusion list", level=1)
    p = doc.add_paragraph("")
    run = p.add_run("File: ")
    run.bold = True
    p.add_run("config/nlc_model.yaml")

    p = doc.add_paragraph("")
    run = p.add_run("Section: ")
    run.bold = True
    p.add_run("dont_update_targets (near the bottom of the file)")

    doc.add_paragraph(
        "Add your new group name as a new list entry under dont_update_targets:"
    )

    add_code_block(doc,
        'dont_update_targets:\n'
        '  - "Margin test"\n'
        '  - "Wm margin split test"\n'
        '  - "Shipping cost added"\n'
        '  - "DSVD test"\n'
        '  - "Increase test"\n'
        '  - "YOUR NEW GROUP"       # <-- add here'
    )

    p = doc.add_paragraph("")
    run = p.add_run("Why: ")
    run.bold = True
    p.add_run(
        "This list tells the low price updates and high price updates methods to "
        "skip SKU-Nodes that belong to a test group. Without this, those SKU-Nodes "
        "would get repriced by the regular update logic AND by the test logic, "
        "causing conflicts."
    )

    p = doc.add_paragraph("")
    run = p.add_run("No code changes needed")
    run.bold = True
    p.add_run(" \u2014 pricing_rules.py reads this list from config at runtime.")

    # ── Step 2 ──
    doc.add_heading("Step 2: Add the pricing method", level=1)
    p = doc.add_paragraph("")
    run = p.add_run("File: ")
    run.bold = True
    p.add_run("src/rules/pricing_rules.py")

    p = doc.add_paragraph("")
    run = p.add_run("Where: ")
    run.bold = True
    p.add_run(
        "Add a new method to the PricingRulesEngine class, before get_new_sku_nodes()."
    )

    doc.add_heading("Method template", level=2)
    doc.add_paragraph(
        "Every test group method follows the same 6-step pattern. Copy this template "
        "and customize the pricing logic in step 2:"
    )

    add_code_block(doc,
        'def get_your_group_updates(self) -> tuple:\n'
        '    """Get updates for SKU-Nodes in [your group name]."""\n'
        '    # 1. Filter to rows tagged with your group\n'
        '    df = self.df_output[\n'
        '        self.df_output["Final target"] == "Your group name"\n'
        '    ].copy()\n'
        '\n'
        '    if len(df) == 0:\n'
        '        logger.info("No [your group] SKU-Nodes found.")\n'
        '        return pd.DataFrame(), pd.DataFrame()\n'
        '\n'
        '    # 2. Compute new Price (CUSTOMIZE THIS PART)\n'
        '    df["Price"] = np.where(\n'
        '        df["Sub-group"] == "Sub A",\n'
        '        df["Final node level cost"],\n'
        '        df["current_nlc_price"],\n'
        '    )\n'
        '\n'
        '    # 3. Price change % and category (always same)\n'
        '    df["Price change %"] = round(\n'
        '        (df["Price"] - df["current_nlc_price"])\n'
        '        / df["current_nlc_price"], 4\n'
        '    )\n'
        '    df["Price change category"] = np.where(\n'
        '        df["Price change %"] < 0, "Decrease",\n'
        '        np.where(df["Price change %"] > 0, "Increase",\n'
        '                 "No change"),\n'
        '    )\n'
        '\n'
        '    # 4. Filter to |delta| >= 1%\n'
        '    df_update = df[\n'
        '        abs(df["Price change %"]) >= self.min_price_change_pct\n'
        '    ].copy()\n'
        '\n'
        '    # 5. DSV format (always same 4 columns)\n'
        '    df_dsv = df_update[\n'
        '        ["Product Code","Identifier","Price","SKU-Node"]\n'
        '    ].rename(columns={\n'
        '        "Product Code": "SKU", "Identifier": "Source"\n'
        '    })\n'
        '\n'
        '    # 6. Tracker format (always same pattern)\n'
        '    df_tracker = df_update[["SKU-Node"]].merge(\n'
        '        self.df_current_tests, on="SKU-Node", how="left"\n'
        '    )\n'
        '    df_tracker["Last price update"] = self.today_str\n'
        '\n'
        '    return df_dsv, df_tracker'
    )

    doc.add_heading("If the method needs external data", level=2)
    doc.add_paragraph(
        "If your group needs data that isn't already in df_output (e.g., DSVD test "
        "needs shipping costs from an external Excel file), add the external data as "
        "a method parameter. The loading happens in pipeline.py (see Step 3)."
    )

    add_code_block(doc,
        'def get_your_group_updates(self,\n'
        '        df_external: pd.DataFrame) -> tuple:\n'
        '    ...\n'
        '    df = df.merge(df_external, how="left",\n'
        '                  on="Identifier")\n'
        '    ...'
    )

    # ── Step 3 ──
    doc.add_heading("Step 3: Wire into the pipeline", level=1)
    p = doc.add_paragraph("")
    run = p.add_run("File: ")
    run.bold = True
    p.add_run("src/pipeline.py")

    p = doc.add_paragraph("")
    run = p.add_run("Three places need changes:")
    run.bold = True

    doc.add_heading("3a. Call the new method (Step 3 \u2014 Pricing Rules)", level=2)
    doc.add_paragraph(
        "Find the block where the other engine.get_* calls are made. Add yours:"
    )

    add_code_block(doc,
        'df_yours_dsv, df_yours_tracker = \\\n'
        '    engine.get_your_group_updates()'
    )

    doc.add_paragraph(
        "If your method needs external data, load it here with a pipeline parameter:"
    )

    add_code_block(doc,
        '# Add parameter: your_data_path: str = None\n'
        'if your_data_path:\n'
        '    df_ext = pd.read_excel(your_data_path, ...)\n'
        '    df_yours_dsv, df_yours_tracker = \\\n'
        '        engine.get_your_group_updates(df_ext)\n'
        'else:\n'
        '    df_yours_dsv, df_yours_tracker = \\\n'
        '        pd.DataFrame(), pd.DataFrame()'
    )

    doc.add_heading("3b. Add to the DSV update list (Step 4 \u2014 Build DSV)", level=2)
    doc.add_paragraph("Find list_dsv_updates = [...] and add your DSV DataFrame:")

    add_code_block(doc,
        'list_dsv_updates = [\n'
        '    df_incr_dsv, df_dsvd_dsv,\n'
        '    df_wm_split_dsv, df_margin_dsv,\n'
        '    df_low_dsv, df_high_dsv,\n'
        '    df_yours_dsv,              # <-- add here\n'
        ']'
    )

    doc.add_heading("3c. Add to the tracker append list (Step 5 \u2014 Tracker)", level=2)
    doc.add_paragraph(
        "Find updater.append_entries([...]) and add your tracker DataFrame:"
    )

    add_code_block(doc,
        'updater.append_entries([\n'
        '    df_new_tracker,\n'
        '    df_low_tracker,\n'
        '    df_high_tracker,\n'
        '    df_wm_split_tracker,\n'
        '    df_margin_tracker,\n'
        '    df_dsvd_tracker,\n'
        '    df_incr_tracker,\n'
        '    df_yours_tracker,          # <-- add here\n'
        '])'
    )

    doc.add_heading("3d. Update logging and Slack", level=2)
    doc.add_paragraph(
        "Add your group to the slack.notify_pricing_rules() dict and the summary "
        "logger lines so it shows up in Slack and logs:"
    )

    add_code_block(doc,
        'slack.notify_pricing_rules({\n'
        '    ...\n'
        '    "Your group": len(df_yours_dsv),\n'
        '})\n'
        '\n'
        'summary = {\n'
        '    ...\n'
        '    "your_group": len(df_yours_dsv),\n'
        '}'
    )

    # ── Step 4 ──
    doc.add_heading("Step 4 (Optional): Add config entries", level=1)
    p = doc.add_paragraph("")
    run = p.add_run("File: ")
    run.bold = True
    p.add_run("config/pricing_rules.yaml")

    doc.add_paragraph(
        "If your group has parameters that might change (thresholds, column names, "
        "sub-group labels), add them under the rules: section:"
    )

    add_code_block(doc,
        'rules:\n'
        '  your_group:\n'
        '    target_label: "Your group name"\n'
        '    description: "What this group does"\n'
        '    some_threshold: 0.06'
    )

    doc.add_paragraph("Then read them in your method:")

    add_code_block(doc,
        'cfg = self._config["rules"]["your_group"]\n'
        'threshold = cfg["some_threshold"]'
    )

    # ── Checklist ──
    doc.add_heading("Checklist", level=1)
    p = doc.add_paragraph("")
    run = p.add_run("When adding a new test group, verify:")
    run.bold = True

    checklist = [
        "Group name added to dont_update_targets in config/nlc_model.yaml",
        "New method added to PricingRulesEngine in src/rules/pricing_rules.py",
        "Method called in src/pipeline.py Step 3 (Pricing Rules)",
        "DSV DataFrame added to list_dsv_updates in src/pipeline.py Step 4",
        "Tracker DataFrame added to updater.append_entries() in src/pipeline.py Step 5",
        "Logging and Slack updated in src/pipeline.py",
        "If external data needed: pipeline parameter added, loading function created",
        "Test with run_pipeline(save=False) to confirm no errors",
    ]
    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")

    # ── Real examples ──
    doc.add_heading("Real Examples", level=1)

    doc.add_heading("Example A: Price Increase test (no external data)", level=2)

    example_a = [
        ("Group name", '"Increase test"'),
        ("Logic", 'If Sub-group is "Increased" or margin < 6%, use Final node level cost; else keep current price'),
        ("External data", "None"),
        ("Method", "get_price_increase_test_updates() in pricing_rules.py"),
        ("Pipeline wiring", "Direct call, no conditional"),
    ]
    add_styled_table(
        doc,
        ["Attribute", "Value"],
        example_a,
        col_widths=[1.5, 5.7],
    )
    doc.add_paragraph("")

    doc.add_heading("Example B: DSVD test (needs external Excel file)", level=2)

    example_b = [
        ("Group name", '"DSVD test"'),
        ("Logic", '"No shipping" gets NLC price; "Shipping cost added" gets NLC + DSVD shipping cost; others keep current'),
        ("External data", "DSVD cost test Excel with shipping costs per node"),
        ("Method", "get_dsvd_test_updates(df_dsvd_test) in pricing_rules.py"),
        ("Pipeline wiring", "dsvd_test_path parameter, _load_dsvd_test_data() helper, conditional call"),
    ]
    add_styled_table(
        doc,
        ["Attribute", "Value"],
        example_b,
        col_widths=[1.5, 5.7],
    )

    # Save
    output_path = "docs/ADDING-TEST-GROUPS.docx"
    doc.save(output_path)
    print(f"Saved to {output_path}")


if __name__ == "__main__":
    build_document()

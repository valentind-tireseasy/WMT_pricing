"""Generate the Correlation Analysis Plan as a .docx file."""

from docx import Document
from docx.shared import Pt

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

doc.add_heading("Correlation Analysis Sub-Project", level=0)
doc.add_paragraph("Implementation Plan", style="Subtitle")

# Metadata
table = doc.add_table(rows=4, cols=2, style="Light List Accent 1")
meta = [
    ("Project", "WalmartPricing NLC Pipeline"),
    ("Sub-project", "SKU-Node Correlation Analysis"),
    ("Date", "2026-03-25"),
    ("Deliverable", "Jupyter notebook: notebooks/correlation_analysis.ipynb"),
]
for i, (k, v) in enumerate(meta):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_paragraph()

# --- 1. Context ---
doc.add_heading("1. Context & Objective", level=1)
doc.add_paragraph(
    "The pricing team needs to understand what drives sales at the SKU-Node level "
    "on Walmart B2B. Currently, pricing decisions (NLC cascade, margin tests, split "
    "tests) are made based on cost/margin thresholds, but there is no systematic "
    "analysis of how price changes, inventory coverage, margin levels, and other "
    "factors correlate with actual sales volume."
)
doc.add_heading("Business Questions", level=2)
for b in [
    "How do NLC price changes affect subsequent sales volume?",
    "What factors (price, margin, inventory, brand, geography) most influence sales?",
    "What margin levels maximize total revenue/profit across SKU-Nodes?",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph(
    "Approach: Build a daily SKU-Node-Date dataset with ~20 features, then run "
    "broad exploratory correlation and statistical analysis."
)

# --- 2. Dataset Specification ---
doc.add_heading("2. Dataset Specification", level=1)

doc.add_heading("2.1 Granularity", level=2)
doc.add_paragraph("One row = one SKU-Node-Date (daily)")
for s in [
    "Date range: Parameterized, default 30 days",
    "SKU filter: Top SKUs accounting for 90% of qty sold over past 2 months",
    "Node filter: All nodes with at least 1 sale for those SKUs",
    "Rollback exclusion: SKUs in active rollbacks (End date > analysis start) excluded",
    "Expected size: ~5,000 SKU-Nodes x 30 days = ~150,000 rows",
]:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("2.2 Core Columns", level=2)
core_t = doc.add_table(rows=19, cols=4, style="Light List Accent 1")
for i, h in enumerate(["#", "Column", "Source", "Computation"]):
    core_t.rows[0].cells[i].text = h
core_data = [
    ("1", "SKU", "Sales query", "Product Code"),
    ("2", "Node", "Sales query", "externalwarehouseid (= Identifier)"),
    ("3", "Date", "Generated", "Daily within analysis window"),
    ("4", "Qty sold", "DW: warehouse.vw_virtual_node_tracker", "Sum of quantity per SKU-Node-Date"),
    ("4b", "Revenue", "DW: warehouse.vw_virtual_node_tracker", "Sum of total_inv_amount per SKU-Node-Date"),
    ("4c", "Profit", "DW: warehouse.vw_virtual_node_tracker", "Sum of profit per SKU-Node-Date"),
    ("5", "Cost to Walmart", "Google Drive DSV files", "Price from DSV; national fallback if no node price"),
    ("6", "Walmart offer price", "DW: pricing_tests.walmart_item_report", "offer_price (status=PUBLISHED)"),
    ("7", "Walmart margin", "Computed", "(offer_price - cost_to_walmart) / offer_price"),
    ("8", "TE Margin", "Computed", "(cost_to_walmart - purchase_price_fet) / cost_to_walmart"),
    ("9", "Brand", "Derived", "First 4 characters of Product Code"),
    ("10", "Size", "DW/Inventory", "Best-effort, not central"),
    ("11", "Town", "Warehouse Addresses CSV", "City of warehouse node"),
    ("12", "State", "Warehouse Addresses CSV", "State of warehouse node"),
    ("13", "Can show inv?", "Inventory (pricing_module)", "1 if inventory available, else 0"),
    ("14", "Min Purchase Price+FET", "Inventory", "Min daily purchase price at node"),
    ("15", "Shipping cost", "Local Excel", "Per-node shipping cost"),
    ("16", "MAP", "DW: pricing_tests.map_prices", "Minimum Advertised Price"),
]
for i, row in enumerate(core_data):
    for j, val in enumerate(row):
        core_t.rows[i + 1].cells[j].text = val

doc.add_heading("2.3 Additional Variables (Claude-suggested, user-approved)", level=2)
extra_t = doc.add_table(rows=5, cols=3, style="Light List Accent 1")
for i, h in enumerate(["#", "Column", "Computation"]):
    extra_t.rows[0].cells[i].text = h
extra_data = [
    ("17", "Day of week", "date.dt.dayofweek (0=Mon, 6=Sun)"),
    ("18", "MAP proximity", "cost_to_walmart / (MAP * 0.95); NaN if no MAP"),
    ("19", "Active nodes per SKU", "Count nodes with can_show_inv=1 per SKU per date"),
    ("20", "Days since last price change", "Days since cost_to_walmart changed for that SKU-Node"),
]
for i, row in enumerate(extra_data):
    for j, val in enumerate(row):
        extra_t.rows[i + 1].cells[j].text = val

doc.add_heading("2.4 Rolling 7-Day Comparison Columns", level=2)
doc.add_paragraph(
    "For each of qty_sold, TE margin, cost_to_walmart, offer_price, walmart_margin:"
)
for r in [
    "{metric}_7d_avg: Mean of prior 7 days (excludes current day)",
    "{metric}_vs_7d: Absolute delta (today minus 7-day average)",
    "{metric}_vs_7d_pct: Percentage change vs 7-day average",
]:
    doc.add_paragraph(r, style="List Bullet")

# --- 3. Data Sources ---
doc.add_heading("3. Data Sources & Loading Strategy", level=1)

doc.add_heading("3.1 Sales Data", level=2)
for s in [
    "Source: warehouse.vw_virtual_node_tracker via DataLoader",
    "Config key: dw_walmart_sales with start_date parameter",
    "Credential mode: new_credentials=false (warehouse.* schema)",
    "Two loads: (1) 2-month lookback for SKU filtering, (2) analysis window + 7-day warmup",
]:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("3.2 DSV Files (Cost to Walmart)", level=2)
for s in [
    "Source: Google Drive folder 1piuawZRpppmoD-Qdkd1IUj3x4rs-LKny",
    "List all files, parse dates, load only those within range (~6-8 files)",
    "For each analysis date, use most recent DSV via merge_asof",
    "National fallback: if no node-specific price, use Source=null row",
    "Rate limiting: 0.5s between API calls (handled by adapter)",
]:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("3.3 Walmart Item Report", level=2)
doc.add_paragraph(
    "Single snapshot for END_DATE. Assumption: offer prices stable over 30 days."
)

doc.add_heading("3.4 Inventory", level=2)
for s in [
    "Source: pricing_module.get_inventory() via lazy import",
    "Sample every 4 days (~8 calls) to manage runtime",
    "Filter: Available >= 4 AND >= Zero Out Threshold",
    "Forward-fill to all analysis dates from nearest inventory date",
    "Runtime: ~5-15 minutes (dominant cost)",
]:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("3.5 Static Reference Data", level=2)
ref_t = doc.add_table(rows=6, cols=3, style="Light List Accent 1")
for i, h in enumerate(["Data", "Source", "Load Method"]):
    ref_t.rows[0].cells[i].text = h
ref_data = [
    ("MAP prices", "pricing_tests.map_prices", 'loader.load("dw_map_prices")'),
    ("Shipping costs", "Local Excel", 'loader.load("shipping_costs_by_node")'),
    ("Warehouse node mapping", "Google Drive", 'loader.load("warehouse_node_mapping")'),
    ("Warehouse addresses", "Local CSV", "pd.read_csv(path)"),
    ("Rollbacks", "Local Excel", 'loader.load("rollbacks") (optional)'),
]
for i, row in enumerate(ref_data):
    for j, val in enumerate(row):
        ref_t.rows[i + 1].cells[j].text = val

# --- 4. Join Chain ---
doc.add_heading("4. Join Chain for City Mapping", level=1)
doc.add_paragraph(
    "Sales.externalwarehouseid = Identifier (node ID)\n"
    "    -> warehouse_node_mapping -> Warehouse Code\n"
    "    -> Warehouse Addresses CSV: Code -> Town, State"
)

# --- 5. Notebook Structure ---
doc.add_heading("5. Notebook Cell Structure", level=1)
sections = [
    ("Section 1: Parameters & Setup", "Configurable parameters, imports, DataLoader init"),
    ("Section 2: SKU Filtering", "Load 2-month sales, select top 90%, exclude rollbacks"),
    ("Section 3: Daily Sales", "Load analysis window + 7-day warmup, aggregate by SKU-Node-Date"),
    ("Section 4: Date Scaffold", "Cross-join SKU-Nodes x dates, left-join sales (fill 0)"),
    ("Section 5: DSV History", "Load DSV files, build price table, merge_asof with national fallback"),
    ("Section 6: Offer Prices", "Load item report snapshot, join to scaffold"),
    ("Section 7: Supporting Data", "MAP, shipping costs, warehouse mapping, addresses"),
    ("Section 8: Inventory", "Lazy import, sampled dates, can_show_inv flag, forward-fill"),
    ("Section 9: Assemble Master DF", "Merge all sources, compute derived columns"),
    ("Section 10: Rolling Comparisons", "7-day rolling avgs, deltas, pct changes, trim warmup"),
    ("Section 11: EDA", "Summary stats, Spearman heatmap, top correlations, scatter plots, distributions"),
    ("Section 12: Statistical Tests", "Mann-Whitney U, margin deciles, inventory vs sales, OLS regression"),
    ("Section 13: Geo & Brand", "Sales by state, top 20 brands, node distribution breadth"),
    ("Section 14: Summary", "Key findings, cleanup"),
]
for title, desc in sections:
    p = doc.add_paragraph()
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

# --- 6. Edge Cases ---
doc.add_heading("6. Edge Cases", level=1)
edge_t = doc.add_table(rows=9, cols=2, style="Light List Accent 1")
edge_t.rows[0].cells[0].text = "Issue"
edge_t.rows[0].cells[1].text = "Handling"
edges = [
    ("No DSV every day", 'Forward-fill via merge_asof(direction="backward")'),
    ("No node-specific DSV price", "Fall back to national price (Source = NaN)"),
    ("Zero-inflated sales", "Spearman corr; log-transform np.log1p(); non-zero subset"),
    ("Rolling window warmup", "Pull 7 extra days before START_DATE, trim after"),
    ("Inventory load perf", "Sample every 4 days, forward-fill between samples"),
    ("Offer price changes", "Single snapshot assumed stable; documented"),
    ("Type mismatches", "Cast all node/identifier/code to str early"),
    ("No MAP for many SKUs", "MAP proximity = NaN; document partial coverage"),
]
for i, (issue, handling) in enumerate(edges):
    edge_t.rows[i + 1].cells[0].text = issue
    edge_t.rows[i + 1].cells[1].text = handling

# --- 7. Dependencies ---
doc.add_heading("7. Dependencies", level=1)
doc.add_paragraph("Already available: pandas, numpy, matplotlib, openpyxl")
doc.add_paragraph("May need installation: seaborn, scipy, statsmodels")

# --- 8. Code Reuse ---
doc.add_heading("8. Existing Code Reused", level=1)
reuse_t = doc.add_table(rows=8, cols=3, style="Light List Accent 1")
for i, h in enumerate(["Component", "File", "What is Reused"]):
    reuse_t.rows[0].cells[i].text = h
reuse = [
    ("DataLoader", "src/data/loader.py", "All data loading, load_dsv_by_date() pattern"),
    ("GoogleAPIAdapter", "src/adapters/google_api_adapter.py", "get_folder_files(), get_file_as_df()"),
    ("DW Adapter", "src/adapters/dw_adapter.py", "SQL query execution"),
    ("Module loader", "src/adapters/module_loader.py", "ensure_modules_path(), load_yaml()"),
    ("Data source configs", "config/data_sources.yaml", "All SQL queries and source definitions"),
    ("NLC model config", "config/nlc_model.yaml", "Inventory thresholds, margin parameters"),
    ("Inventory pattern", "src/models/nlc_model.py", "_load_inventory() method as template"),
]
for i, row in enumerate(reuse):
    for j, val in enumerate(row):
        reuse_t.rows[i + 1].cells[j].text = val

# --- 9. Verification ---
doc.add_heading("9. Verification Checklist", level=1)
for c in [
    "Run notebook with ANALYSIS_DAYS=7 first (smaller, faster)",
    "Check scaffold shape matches n_sku_nodes x n_days",
    "Verify no unexpected NaN patterns in cost_to_walmart, offer_price",
    "Confirm rolling averages are clean after warmup trim",
    "Spot-check a few SKU-Nodes against actual DSV files",
    "Correlation matrix renders correctly",
    "Top correlations make business sense",
    "Statistical tests produce interpretable p-values",
]:
    doc.add_paragraph(c, style="List Bullet")

doc.save("docs/correlation-analysis-plan.docx")
print("Done: docs/correlation-analysis-plan.docx")
